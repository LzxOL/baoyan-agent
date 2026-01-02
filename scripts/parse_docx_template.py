#!/usr/bin/env python3
"""
Parse a DOCX template and extract:
- embedded image shapes (logo) extents (cx, cy in EMU) and derived pixel size
- text placeholders to be replaced (heuristics: right-side of '：' or ':' or runs with 'XXX'/'____')

Usage:
  python scripts/parse_docx_template.py /path/to/template.docx /path/to/output.json

Dependencies:
  pip install python-docx lxml

Notes:
 - This script uses python-docx's oxml to inspect wp:inline drawing elements to get extent (cx/cy).
 - Position information in DOCX is limited; this extracts size (extent) and relationship id to the image part.
 - Placeholder detection uses simple heuristics; inspect the JSON output and adjust rules if needed.
"""
from __future__ import annotations

import json
import math
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document

# Namespaces used for XPath
NSMAP = {
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}


def emu_to_px(emu: int, dpi: int = 96) -> float:
    """
    Convert EMU (English Metric Unit used in DOCX) to pixels.
    1 inch = 914400 EMU. px = (emu / 914400) * dpi
    Default dpi=96 is common for screens.
    """
    return (emu / 914400.0) * dpi


@dataclass
class ImageInfo:
    rId: str
    partname: str
    cx: int
    cy: int
    width_px: float
    height_px: float


@dataclass
class PlaceholderInfo:
    paragraph_index: int
    label: str
    placeholder_text: str
    run_index: int
    font_name: Optional[str]
    font_size_pt: Optional[float]


def extract_images_info(doc: Document) -> List[Dict]:
    images: List[ImageInfo] = []
    # search for inline drawings in paragraphs' runs
    for p_idx, paragraph in enumerate(doc.paragraphs):
        for r_idx, run in enumerate(paragraph.runs):
            # find wp:inline or wp:anchor within this run's r element using expanded QNames
            inline_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            anchor_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            nodes = inline_nodes + anchor_nodes
            for node in nodes:
                # extent element holds cx, cy in EMU
                extent = node.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                cx = cy = None
                if extent:
                    try:
                        cx = int(extent[0].get('cx'))
                        cy = int(extent[0].get('cy'))
                    except Exception:
                        cx = cy = None

                # blip embed holds r:embed which maps to relationship id
                blip = node.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                rId = None
                partname = ''
                if blip:
                    rId = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        # rel._target is the part; expose partname (e.g. /word/media/image1.png)
                        try:
                            partname = rel.target_part.partname if hasattr(rel, 'target_part') else str(rel.target_ref)
                        except Exception:
                            partname = str(getattr(rel, 'target_ref', ''))

                if rId and cx and cy:
                    img = ImageInfo(
                        rId=rId,
                        partname=str(partname),
                        cx=cx,
                        cy=cy,
                        width_px=round(emu_to_px(cx), 2),
                        height_px=round(emu_to_px(cy), 2),
                    )
                    images.append(img)
    return [asdict(i) for i in images]


def detect_placeholders(doc: Document) -> List[Dict]:
    """
    Heuristics:
    - If a paragraph contains '：' or ':' treat left of first colon as label, right as placeholder
    - If runs contain 'XXX' or '____' treat those runs as placeholders
    """
    placeholders: List[PlaceholderInfo] = []
    for p_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text or ''
        # full-width colon
        if '：' in text or ':' in text:
            # split at first colon (prefer full-width)
            if '：' in text:
                left, right = text.split('：', 1)
            else:
                left, right = text.split(':', 1)
            placeholder_text = right.strip()
            label = left.strip()
            # find the run index that contains the placeholder_text (or the last run)
            run_index = 0
            font_name = None
            font_size_pt = None
            found = False
            for r_idx, run in enumerate(paragraph.runs):
                run_text = run.text or ''
                if placeholder_text and placeholder_text in run_text:
                    run_index = r_idx
                    font_name = run.font.name
                    font_size_pt = run.font.size.pt if run.font.size else None
                    found = True
                    break
            if not found:
                # fallback: use last run's font info
                if paragraph.runs:
                    run_index = len(paragraph.runs) - 1
                    last = paragraph.runs[-1]
                    font_name = last.font.name
                    font_size_pt = last.font.size.pt if last.font.size else None

            placeholders.append(
                PlaceholderInfo(
                    paragraph_index=p_idx,
                    label=label,
                    placeholder_text=placeholder_text,
                    run_index=run_index,
                    font_name=font_name,
                    font_size_pt=font_size_pt,
                )
            )
            continue

        # fallback: look for common placeholder tokens in runs
        for r_idx, run in enumerate(paragraph.runs):
            t = (run.text or '').strip()
            if not t:
                continue
            if 'XXX' in t or '____' in t or t.startswith('[') and t.endswith(']'):
                font_name = run.font.name
                font_size_pt = run.font.size.pt if run.font.size else None
                placeholders.append(
                    PlaceholderInfo(
                        paragraph_index=p_idx,
                        label='',
                        placeholder_text=t,
                        run_index=r_idx,
                        font_name=font_name,
                        font_size_pt=font_size_pt,
                    )
                )
    return [asdict(p) for p in placeholders]


def main(argv):
    if len(argv) < 2:
        print("Usage: parse_docx_template.py /path/to/template.docx [/path/to/output.json]")
        return 2
    docx_path = Path(argv[1])
    if not docx_path.exists():
        print("File not found:", docx_path)
        return 2
    out_path = Path(argv[2]) if len(argv) > 2 else Path('scripts/template_parsed.json')

    doc = Document(str(docx_path))

    images = extract_images_info(doc)
    placeholders = detect_placeholders(doc)

    result = {
        'source': str(docx_path),
        'images': images,
        'placeholders': placeholders,
        'summary': {
            'image_count': len(images),
            'placeholder_count': len(placeholders),
        },
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2))
    print("Wrote parsed template to", out_path)
    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv))


