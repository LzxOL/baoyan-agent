#!/usr/bin/env python3
"""
Generate a cover PDF/DOCX for a given school using a DOCX template and local logo images.

Usage:
  python scripts/generate_school_cover.py \
    --template "/home/root1/推免生证明材料（模板）.docx" \
    --logos "/home/root1/下载/中国所有大学校徽图片-200px-jpgs" \
    --school "清华大学" \
    --output "/home/root1/output/tsinghua-cover.pdf"

Behavior:
 - Finds a logo file in the logos directory matching the school name (case-insensitive substring).
 - Replaces the first image in the template with that logo (preserving size).
 - Replaces placeholder fields detected by simple heuristics:
     - "标签：占位" 形式会把右侧替换为 provided values (if present).
     - runs containing 'XXX' or '____' will be replaced in-order with provided mapping.
 - Saves a modified DOCX and attempts to convert it to PDF using LibreOffice (soffice).

Dependencies:
  pip install python-docx lxml
  LibreOffice (optional, for PDF conversion)
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Loaded template spec (can be overridden via --spec)
TEMPLATE_SPEC: Dict = {}

def emu_to_px(emu: int, dpi: int = 300) -> float:
    """
    Convert EMU (English Metric Unit) to pixels.
    1 inch = 914400 EMU, default DPI is 300 for high-quality printing.
    """
    inches = emu / 914400.0
    return inches * dpi

def find_logo_file(logos_dir: Path, school_name: str) -> Optional[Path]:
    # First try explicit mapping from spec-generated mapping file
    mapped = None
    try:
        logo_map_path = Path(TEMPLATE_SPEC.get('logo_mapping', Path(__file__).parent / 'logo_mapping.json'))
        if logo_map_path.exists():
            mapped_json = json.loads(logo_map_path.read_text(encoding='utf-8'))
            # exact match
            if school_name in mapped_json:
                candidate = logos_dir / mapped_json[school_name]
                if candidate.exists():
                    return candidate
            # try simplified match keys
            name_low = school_name.lower()
            for k, v in mapped_json.items():
                if k.lower() == name_low:
                    candidate = logos_dir / v
                    if candidate.exists():
                        return candidate
    except Exception:
        pass
    name_low = school_name.lower()
    candidates: List[Path] = []
    for p in logos_dir.iterdir():
        if not p.is_file():
            continue
        fname = p.name.lower()
        if name_low in fname:
            candidates.append(p)
    if not candidates:
        # try fuzzy: split school_name into tokens
        tokens = re.split(r'[\s\-]+', school_name)
        for p in logos_dir.iterdir():
            fname = p.name.lower()
            if all(tok.lower() in fname for tok in tokens if tok):
                candidates.append(p)
    if not candidates:
        return None
    # prefer svg, then png, then jpg
    def score(x: Path):
        ext = x.suffix.lower()
        if ext == '.svg':
            return 0
        if ext == '.png':
            return 1
        if ext in ('.jpg', '.jpeg'):
            return 2
        return 3
    candidates.sort(key=lambda x: (score(x), x.name))
    return candidates[0]

def replace_first_image_with_logo(doc: Document, logo_path: Path) -> bool:
    """
    Replace the blob of the first image relationship found in the document with the logo bytes.
    Returns True on success.
    """
    from docx.shared import Emu

    # First attempt: replace bytes on first image part found in package
    try:
        with open(logo_path, 'rb') as f:
            blob = f.read()
        for rel in list(doc.part.rels.values()):
            try:
                part = getattr(rel, 'target_part', None)
                if not part:
                    continue
                ctype = getattr(part, 'content_type', '') or ''
                if ctype.startswith('image/'):
                    try:
                        part._blob = blob
                        return True
                    except Exception:
                        try:
                            part.blob = blob
                            return True
                        except Exception:
                            continue
            except Exception:
                continue
    except Exception:
        pass

    # Fallback: search for first image run; replace by inserting picture in same run
    for p in doc.paragraphs:
        for run in p.runs:
            # find drawing or pict elements inside run xml
            drawing_nodes = []
            for child in list(run._r):
                tag = getattr(child, 'tag', '')
                if tag is None:
                    continue
                if tag.endswith('}drawing') or tag.endswith('}pict') or 'drawing' in tag:
                    drawing_nodes.append(child)

            if not drawing_nodes:
                # also check deeper inline/anchor nodes
                inline_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                anchor_nodes = run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                nodes = inline_nodes + anchor_nodes
                if not nodes:
                    continue
                node = nodes[0]
                # try to get extent cx/cy
                extent = node.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                cx = cy = None
                if extent:
                    try:
                        cx = int(extent[0].get('cx'))
                        cy = int(extent[0].get('cy'))
                    except Exception:
                        cx = cy = None
                # remove the existing drawing elements under run
                for child in list(run._r):
                    try:
                        tag = getattr(child, 'tag', '')
                        if tag and (tag.endswith('}drawing') or tag.endswith('}pict') or 'drawing' in tag):
                            run._r.remove(child)
                    except Exception:
                        pass

                # insert new picture in this run with same size if available
                try:
                    # If logo is SVG, rasterize it to temp PNG with desired pixel size
                    tmp_to_remove = None
                    if logo_path.suffix.lower() == '.svg':
                        try:
                            import cairosvg
                            # convert EMU to px for rasterization (1 inch = 914400 EMU)
                            # Use higher DPI (300) and scale factor (3x) for crisp logos when zoomed in PDF
                            scale_factor = 3.0  # 3x scale for high-quality rendering
                            px_w = round(emu_to_px(cx, dpi=300) * scale_factor) if cx else None
                            px_h = round(emu_to_px(cy, dpi=300) * scale_factor) if cy else None
                            tmp_png = Path.cwd() / f".tmp_logo_{os.getpid()}.png"
                            if px_w and px_h:
                                cairosvg.svg2png(url=str(logo_path), write_to=str(tmp_png), output_width=px_w, output_height=px_h)
                            elif px_w:
                                cairosvg.svg2png(url=str(logo_path), write_to=str(tmp_png), output_width=px_w)
                            else:
                                cairosvg.svg2png(url=str(logo_path), write_to=str(tmp_png))
                            logo_to_use = tmp_png
                            tmp_to_remove = tmp_png
                        except Exception:
                            # fallback to using original path
                            logo_to_use = logo_path
                    else:
                        logo_to_use = logo_path

                    if cx and cy:
                        run.add_picture(str(logo_to_use), width=Emu(cx), height=Emu(cy))
                    else:
                        run.add_picture(str(logo_to_use))
                    if tmp_to_remove and tmp_to_remove.exists():
                        try:
                            tmp_to_remove.unlink()
                        except Exception:
                            pass
                    return True
                except Exception:
                    # insertion failed; try next run
                    continue
    return False

def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> int:
    """
    Replace placeholders in 'label: value' style and simple tokens.
    mapping: { '学生姓名': '张三', '申请专业': '计算机' }
    Returns number of replacements made.
    """
    replaced = 0

    # Helper to copy font properties from one run to another
    def copy_font_props(src_run, dst_run):
        if not hasattr(src_run, 'font') or not hasattr(dst_run, 'font'):
            return
        try:
            dst_run.font.size = src_run.font.size
            dst_run.font.name = src_run.font.name
            dst_run.font.bold = src_run.font.bold
            dst_run.font.italic = src_run.font.italic
            dst_run.font.underline = src_run.font.underline
        except Exception:
            pass

    # Load keys priority and placeholder chars from spec if provided
    keys_priority = TEMPLATE_SPEC.get('keys_priority', ['学生姓名', '申请专业', '本科院校', '毕业专业', '联系方式', '邮箱'])
    placeholder_chars = TEMPLATE_SPEC.get('placeholder_chars', ['×', 'X'])
    placeholder_set = set(placeholder_chars)

    def _align_from_str(s: str):
        if not s:
            return None
        s = s.lower()
        if s in ('left', 'l'):
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        if s in ('right', 'r'):
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        if s in ('center', 'centre', 'c'):
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        return None

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text or ''
        modified_run_ids = set()
        def clear_adjacent_placeholder_runs(runs, idx):
            # clear following runs that contain only placeholder characters (e.g. × or X)
            for j in range(idx + 1, len(runs)):
                text = runs[j].text or ''
                if text.strip() and all(ch in placeholder_set for ch in text.strip()):
                    runs[j].text = ''
                    modified_run_ids.add(id(runs[j]))
                else:
                    break

        # 1) Handle label: value lines (e.g. "学  生  姓  名  ：   ×××")
        if '：' in paragraph_text or ':' in paragraph_text:
            if '：' in paragraph_text:
                left, right = paragraph_text.split('：', 1)
                sep = '：'
            else:
                left, right = paragraph_text.split(':', 1)
                sep = ':'
            label = left.strip()

            # Find which key this label corresponds to by checking known labels in runs
            matched_key = None
            # simple heuristics: check runs for label tokens
            for key in keys_priority:
                if key in mapping:
                    # check a few common Chinese label forms
                    if key == '学生姓名' and ('学' in label and '名' in label):
                        matched_key = key
                        break
                    if key == '申请专业' and '专' in label and '业' in label:
                        matched_key = key
                        break
                    if key == '本科院校' and ('院' in label and '校' in label):
                        matched_key = key
                        break
                    if key == '毕业专业' and '毕' in label and '业' in label:
                        matched_key = key
                        break
                    if key == '联系方式' and '联' in label and '系' in label:
                        matched_key = key
                        break
                    if key == '邮箱' and '邮' in label:
                        matched_key = key
                        break

            # If matched, replace only the run(s) that contain × or the right-hand side
            if matched_key:
                # create a 1x2 borderless table at the paragraph position:
                # left cell: label+sep, right cell: value
                runs = paragraph.runs
                placeholder_run = None
                for r in runs:
                    rt = r.text or ''
                    if any(pc in rt for pc in placeholder_chars) or right.strip() in rt:
                        placeholder_run = r
                        break

                # Build table and insert after paragraph
                table = doc.add_table(rows=1, cols=2)
                # set column widths and vertical alignment for nicer layout
                try:
                    table.allow_autofit = False
                    # left/right column widths from template spec if provided
                    left_w = TEMPLATE_SPEC.get('table', {}).get('left_col_width_in', 2.2)
                    right_w = TEMPLATE_SPEC.get('table', {}).get('right_col_width_in', 4.0)
                    table.columns[0].width = Inches(left_w)
                    table.columns[1].width = Inches(right_w)
                    # set vertical alignment center for both cells
                    for cell in (table.cell(0, 0), table.cell(0, 1)):
                        tc = cell._tc
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            tc.insert(0, tcPr)
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)
                except Exception:
                    pass
                # remove table borders to make it borderless
                try:
                    tbl = table._tbl
                    tblPr = tbl.find(qn('w:tblPr'))
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    tblBorders = OxmlElement('w:tblBorders')
                    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'nil')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)
                except Exception:
                    pass
                # left cell
                left_cell = table.cell(0, 0)
                right_cell = table.cell(0, 1)

                # fill left cell with label and separator, preserving simple formatting
                left_para = left_cell.paragraphs[0]
                left_para.text = left + sep
                # right-align label so its visual center aligns with cell middle (configurable)
                left_align = TEMPLATE_SPEC.get('table', {}).get('left_cell_alignment', 'right')
                try:
                    la = _align_from_str(left_align)
                    if la is not None:
                        left_para.alignment = la
                except Exception:
                    pass
                if runs:
                    try:
                        src_run = runs[0]
                        if left_para.runs:
                            copy_font_props(src_run, left_para.runs[0])
                    except Exception:
                        pass

                # fill right cell with mapped value, preserving placeholder run font if possible
                right_para = right_cell.paragraphs[0]
                right_para.text = mapping[matched_key]
                # align value inside right cell per spec (e.g., center)
                right_align = TEMPLATE_SPEC.get('table', {}).get('right_cell_alignment', 'center')
                try:
                    ra = _align_from_str(right_align)
                    if ra is not None:
                        right_para.alignment = ra
                except Exception:
                    pass
                if placeholder_run:
                    try:
                        if right_para.runs:
                            copy_font_props(placeholder_run, right_para.runs[0])
                    except Exception:
                        pass

                # insert table after current paragraph and remove the paragraph
                try:
                    paragraph._p.addnext(table._tbl)
                    paragraph._p.getparent().remove(paragraph._p)
                except Exception:
                    # fallback: if insertion fails, just replace in-place
                    for i, run in enumerate(runs):
                        rt = run.text or ''
                    if any(pc in rt for pc in placeholder_chars) or right.strip() in rt:
                            leading_ws = rt[:len(rt) - len(rt.lstrip())] if rt else ''
                            run.text = leading_ws + mapping[matched_key]
                            modified_run_ids.add(id(run))
                            replaced += 1
                            clear_adjacent_placeholder_runs(runs, i)
                            break
                replaced += 1
                continue

        # 2) Handle header like "××大学××学院" where runs are split (××, 大学, ××, 学院)
        if '大学' in paragraph_text and '学院' in paragraph_text and any(pc in paragraph_text for pc in placeholder_chars):
            # find runs for the pattern and replace only the × runs
            runs = paragraph.runs
            # find index of first '大学' run and '学院' run if present
            uni_idx = None
            coll_idx = None
            for i, r in enumerate(runs):
                if r.text == '大学':
                    uni_idx = i
                if r.text == '学院':
                    coll_idx = i
            # Replace run before '大学' (if contains ×) with university short name (no trailing '大学')
            if uni_idx is not None and uni_idx > 0:
                candidate = runs[uni_idx - 1]
                if any(pc in (candidate.text or '') for pc in placeholder_chars) and id(candidate) not in modified_run_ids:
                    uni_full = mapping.get('本科院校', mapping.get('学校', '清华大学'))
                    # remove trailing '大学' to avoid duplication because the next run is '大学'
                    uni_short = uni_full[:-2] if uni_full.endswith('大学') else uni_full
                    candidate.text = uni_short
                    modified_run_ids.add(id(candidate))
                    replaced += 1
                    # clear any adjacent placeholder runs (e.g., duplicates)
                    clear_adjacent_placeholder_runs(runs, uni_idx - 1)
            # Replace run before '学院' (if contains ×) with department derived from major
            if coll_idx is not None and coll_idx > 0:
                candidate = runs[coll_idx - 1]
                if any(pc in (candidate.text or '') for pc in placeholder_chars) and id(candidate) not in modified_run_ids:
                    dept = mapping.get('申请专业', mapping.get('毕业专业', '计算机科学与技术'))
                    # keep only the department name (append nothing, since next run is '学院')
                    candidate.text = dept
                    modified_run_ids.add(id(candidate))
                    replaced += 1
                    clear_adjacent_placeholder_runs(runs, coll_idx - 1)

        # 3) Generic run-level replacements for any remaining × placeholders
        for i, run in enumerate(paragraph.runs):
            rt = run.text or ''
            if any(pc in rt for pc in placeholder_chars):
                if id(run) in modified_run_ids:
                    continue
                # find best key by paragraph context
                chosen_key = None
                for key in keys_priority:
                    if key in mapping:
                        # context clues
                        if key == '学生姓名' and '姓' in paragraph_text and '名' in paragraph_text:
                            chosen_key = key
                            break
                        if key == '申请专业' and '申' in paragraph_text:
                            chosen_key = key
                            break
                        if key == '本科院校' and '院' in paragraph_text and '校' in paragraph_text:
                            chosen_key = key
                            break
                        if key == '毕业专业' and '毕' in paragraph_text:
                            chosen_key = key
                            break
                        if key == '联系方式' and '联' in paragraph_text:
                            chosen_key = key
                            break
                        if key == '邮箱' and '邮' in paragraph_text:
                            chosen_key = key
                            break
                if not chosen_key:
                    # fallback: pick first available mapping
                    chosen_key = next(iter(mapping.keys()))
                # preserve leading whitespace, then set mapping value
                leading_ws = rt[:len(rt) - len(rt.lstrip())] if rt else ''
                run.text = leading_ws + mapping.get(chosen_key, '')
                modified_run_ids.add(id(run))
                replaced += 1
                # clear following placeholder-only runs to avoid duplication
                clear_adjacent_placeholder_runs(paragraph.runs, i)

    return replaced

def convert_docx_to_pdf(docx_path: Path, out_pdf: Path) -> bool:
    """
    Try to convert using LibreOffice (soffice). Returns True if pdf created.
    """
    if not shutil.which('soffice'):
        print("LibreOffice (soffice) not found; skipping PDF conversion.")
        return False
    outdir = out_pdf.parent
    cmd = [
        'soffice', '--headless', '--convert-to', 'pdf', str(docx_path),
        '--outdir', str(outdir)
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        # soffice names pdf same basename
        generated = outdir / (docx_path.stem + '.pdf')
        if generated.exists():
            generated.rename(out_pdf)
            return True
    except subprocess.CalledProcessError as e:
        print("LibreOffice conversion failed:", e)
    return False

def main(argv: List[str]):
    parser = argparse.ArgumentParser()
    parser.add_argument('--template', required=True, help='Path to DOCX template')
    parser.add_argument('--logos', required=True, help='Path to logos directory')
    parser.add_argument('--school', required=True, help='School name to match logo')
    parser.add_argument('--output', required=True, help='Output PDF path (or .docx)')
    parser.add_argument('--fields', help='JSON string or file with mapping of placeholders to values')
    parser.add_argument('--spec', help='Path to template spec JSON (optional)')
    args = parser.parse_args(argv)

    template = Path(args.template)
    logos_dir = Path(args.logos)
    school = args.school
    output = Path(args.output)
    spec_path = None
    if getattr(args, 'spec', None):
        spec_path = Path(args.spec)
    else:
        spec_path = Path(__file__).parent / 'template_spec.json'
    try:
        if spec_path and spec_path.exists():
            global TEMPLATE_SPEC
            TEMPLATE_SPEC = json.loads(spec_path.read_text(encoding='utf-8'))
            print("Loaded template spec from", spec_path)
    except Exception as e:
        print("Failed to load template spec:", e)

    if not template.exists():
        print("Template not found:", template)
        return 2
    if not logos_dir.exists():
        print("Logos dir not found:", logos_dir)
        return 2

    mapping: Dict[str, str] = {}
    if args.fields:
        try:
            # allow passing a path to json file
            fpath = Path(args.fields)
            if fpath.exists():
                mapping = json.loads(fpath.read_text(encoding='utf-8'))
            else:
                mapping = json.loads(args.fields)
        except Exception as e:
            print("Failed to parse fields mapping:", e)
            return 2

    # If no fields provided or empty, use default test data from spec (if available)
    if not mapping:
        mapping = TEMPLATE_SPEC.get('defaults', {
            "学生姓名": "王小明",
            "申请专业": "计算机科学与技术",
            "本科院校": "北京大学",
            "毕业专业": "软件工程",
            "联系方式": "138-0000-0000",
            "邮箱": "wangxiaoming@pku.edu.cn"
        })
        print("Using default test data for placeholders")

    logo_file = find_logo_file(logos_dir, school)
    if not logo_file:
        print("Logo file not found for school:", school)
        # continue but warn
    else:
        print("Using logo:", logo_file)

    doc = Document(str(template))

    if logo_file:
        ok = replace_first_image_with_logo(doc, logo_file)
        print("Logo replace:", ok)

    replaced = replace_placeholders(doc, mapping)
    print("Placeholders replaced:", replaced)

    out_docx = output.with_suffix('.docx') if output.suffix.lower() != '.docx' else output
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print("Saved generated docx to", out_docx)

    if output.suffix.lower() == '.pdf':
        ok = convert_docx_to_pdf(out_docx, output)
        if ok:
            print("Saved PDF to", output)
        else:
            print("PDF not generated; docx saved at", out_docx)

    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv[1:]))


